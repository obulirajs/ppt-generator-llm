"""
Document Parser Module
Extracts text from various document formats for PPT generation
Supports: .txt, .docx, .pdf files and direct text input
"""

import os
import uuid
from werkzeug.utils import secure_filename
from docx import Document
import PyPDF2


# Configuration Constants
ALLOWED_EXTENSIONS = {'txt', 'docx', 'pdf'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB in bytes
MIN_CONTENT_LENGTH = 50  # Minimum characters for meaningful content
MAX_CONTENT_LENGTH = 100000  # Maximum characters to process


def allowed_file(filename):
    """
    Check if file has an allowed extension
    
    Args:
        filename (str): Name of the file
        
    Returns:
        bool: True if extension is allowed, False otherwise
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_file_extension(filename):
    """
    Extract file extension from filename
    
    Args:
        filename (str): Name of the file
        
    Returns:
        str: File extension in lowercase (without dot)
    """
    if '.' in filename:
        return filename.rsplit('.', 1)[1].lower()
    return ''


def validate_file(file):
    """
    Validate uploaded file for size and type
    
    Args:
        file: FileStorage object from Flask
        
    Returns:
        tuple: (is_valid: bool, error_message: str or None)
    """
    # Check if file exists
    if not file or file.filename == '':
        return False, "No file provided"
    
    # Check file extension
    if not allowed_file(file.filename):
        return False, f"Invalid file type. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
    
    # Check file size (seek to end to get size)
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)  # Reset file pointer to beginning
    
    if file_size > MAX_FILE_SIZE:
        size_mb = MAX_FILE_SIZE / (1024 * 1024)
        return False, f"File too large. Maximum size: {size_mb}MB"
    
    if file_size == 0:
        return False, "File is empty"
    
    return True, None


def validate_content(text):
    """
    Validate extracted text content
    
    Args:
        text (str): Extracted text content
        
    Returns:
        tuple: (is_valid: bool, error_message: str or None, processed_text: str)
    """
    if not text or not text.strip():
        return False, "No text content found in the document", ""
    
    text = text.strip()
    
    # Check minimum length
    if len(text) < MIN_CONTENT_LENGTH:
        return False, f"Content too short. Minimum {MIN_CONTENT_LENGTH} characters required", text
    
    # Truncate if too long
    if len(text) > MAX_CONTENT_LENGTH:
        text = text[:MAX_CONTENT_LENGTH]
        warning = f"Content truncated to {MAX_CONTENT_LENGTH} characters"
        return True, warning, text
    
    return True, None, text


def extract_from_txt(file_path):
    """
    Extract text from plain text file
    
    Args:
        file_path (str): Path to the .txt file
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If file cannot be read
    """
    try:
        # Try UTF-8 first
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Fallback to latin-1 if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Failed to read text file: {str(e)}")
    except Exception as e:
        raise Exception(f"Failed to read text file: {str(e)}")


def extract_from_docx(file_path):
    """
    Extract text from Word document (.docx)
    
    Args:
        file_path (str): Path to the .docx file
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If file cannot be read or parsed
    """
    try:
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Extract text from tables (optional but useful)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
    
    except Exception as e:
        raise Exception(f"Failed to read Word document: {str(e)}")


def extract_from_pdf(file_path):
    """
    Extract text from PDF file
    
    Args:
        file_path (str): Path to the .pdf file
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If file cannot be read or parsed
    """
    try:
        full_text = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Check if PDF has pages
            if len(pdf_reader.pages) == 0:
                raise Exception("PDF file has no pages")
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                if text.strip():
                    full_text.append(text)
        
        extracted = '\n'.join(full_text)
        
        if not extracted.strip():
            raise Exception("No text could be extracted from PDF (might be image-based)")
        
        return extracted
    
    except Exception as e:
        raise Exception(f"Failed to read PDF file: {str(e)}")


def extract_text_from_file(file_path, file_type):
    """
    Main function to extract text from any supported file type
    
    Args:
        file_path (str): Path to the file
        file_type (str): File extension (txt, docx, pdf)
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If file type is unsupported or extraction fails
    """
    file_type = file_type.lower()
    
    # Route to appropriate extraction function
    if file_type == 'txt':
        return extract_from_txt(file_path)
    elif file_type == 'docx':
        return extract_from_docx(file_path)
    elif file_type == 'pdf':
        return extract_from_pdf(file_path)
    else:
        raise Exception(f"Unsupported file type: {file_type}")


def save_uploaded_file(file, upload_folder):
    """
    Save uploaded file to temporary directory with unique filename
    
    Args:
        file: FileStorage object from Flask
        upload_folder (str): Path to upload directory
        
    Returns:
        tuple: (file_path: str, original_filename: str)
        
    Raises:
        Exception: If file cannot be saved
    """
    try:
        # Create upload folder if it doesn't exist
        os.makedirs(upload_folder, exist_ok=True)
        
        # Secure the filename
        original_filename = secure_filename(file.filename)
        
        # Generate unique filename to avoid conflicts
        unique_id = str(uuid.uuid4())
        file_extension = get_file_extension(original_filename)
        unique_filename = f"{unique_id}.{file_extension}"
        
        # Full path for saving
        file_path = os.path.join(upload_folder, unique_filename)
        
        # Save the file
        file.save(file_path)
        
        return file_path, original_filename
    
    except Exception as e:
        raise Exception(f"Failed to save uploaded file: {str(e)}")


def process_direct_text(text):
    """
    Process text that was directly pasted/typed by user
    
    Args:
        text (str): Direct text input from user
        
    Returns:
        str: Processed text
        
    Raises:
        Exception: If text is invalid
    """
    if not text or not isinstance(text, str):
        raise Exception("No text content provided")
    
    text = text.strip()
    
    if len(text) < MIN_CONTENT_LENGTH:
        raise Exception(f"Text too short. Minimum {MIN_CONTENT_LENGTH} characters required")
    
    # Truncate if too long
    if len(text) > MAX_CONTENT_LENGTH:
        text = text[:MAX_CONTENT_LENGTH]
    
    return text


def cleanup_temp_file(file_path):
    """
    Delete temporary file after processing
    
    Args:
        file_path (str): Path to temporary file
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        # Log error but don't raise - file cleanup is not critical
        print(f"Warning: Could not delete temporary file {file_path}: {str(e)}")